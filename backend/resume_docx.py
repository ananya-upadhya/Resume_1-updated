"""
ATS-Compliant DOCX Generator
Matches the exact data shape from ResumeBuilder.jsx:
  - personal: { name, title, email, phone, location, linkedin, github, website }
  - summary:  { text }
  - experience: [{ id, company, role, start, end, current, bullets (newline string) }]
  - education:  [{ id, institution, degree, field, start, end, gpa }]
  - skills:     [str, ...]   ← flat array of strings
  - projects:   [{ id, name, tech, url, bullets (newline string) }]
  - certifications: [{ id, name, issuer, date }]

ATS rules obeyed:
  - No tables, text boxes, columns or images
  - All content in document body (not header/footer)
  - Standard fonts (Calibri)
  - Real OOXML bullet lists (no unicode symbols)
  - Correct XML element ordering (pBdr/tabs inserted at position 0)
  - w:zoom w:percent attribute present
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import zipfile, shutil, os, tempfile, re


# Colours
NAVY   = RGBColor(0x1F, 0x4E, 0x79)
BLACK  = RGBColor(0x11, 0x18, 0x27)
GREY   = RGBColor(0x4B, 0x55, 0x63)
BLUE   = RGBColor(0x3B, 0x82, 0xF6)


# ── helpers ───────────────────────────────────────────────────────────────────

def _parse_bullets(text: str) -> list:
    return [l.strip().lstrip("-•▸* ") for l in (text or "").split("\n") if l.strip()]

def _fmt_month(m: str) -> str:
    """'2024-06' → 'Jun 2024'"""
    if not m: return ""
    parts = m.split("-")
    if len(parts) == 2:
        months = ["Jan","Feb","Mar","Apr","May","Jun","Jul","Aug","Sep","Oct","Nov","Dec"]
        try: return f"{months[int(parts[1])-1]} {parts[0]}"
        except: pass
    return m

def _run(para, text, bold=False, italic=False, size=10, color=BLACK):
    r = para.add_run(text)
    r.bold = bold; r.italic = italic
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return r

def _para(doc, space_before=0, space_after=0, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    return p

def _section_heading(doc, title: str):
    p = _para(doc, space_before=7, space_after=2)
    r = p.add_run(title.upper())
    r.bold = True; r.font.name = "Calibri"
    r.font.size = Pt(10.5); r.font.color.rgb = NAVY
    # Bottom border — insert at pos 0 so XML ordering is correct
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "6")
    bot.set(qn("w:space"), "1");    bot.set(qn("w:color"), "3B82F6")
    pBdr.append(bot)
    pPr.insert(0, pBdr)

def _two_col(doc, left: str, right: str, left_bold=False, left_size=10,
             right_size=9, left_color=BLACK, right_color=GREY):
    p = _para(doc, space_before=1, space_after=1)
    # Tab stop at right edge — insert at pos 0 for XML ordering
    pPr = p._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab  = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right"); tab.set(qn("w:pos"), "9360")
    tabs.append(tab); pPr.insert(0, tabs)
    _run(p, left,  bold=left_bold,  size=left_size,  color=left_color)
    p.add_run("\t").font.name = "Calibri"
    _run(p, right, bold=False,      size=right_size, color=right_color)

def _bullet(doc, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent  = Inches(0.2)
    r = p.add_run(text)
    r.font.name = "Calibri"; r.font.size = Pt(9.5); r.font.color.rgb = BLACK

def _fix_docx(path: str):
    """Post-process: fix w:zoom missing w:percent attribute."""
    tmp_dir = tempfile.mkdtemp()
    try:
        with zipfile.ZipFile(path) as z:
            z.extractall(tmp_dir)
        sx_path = os.path.join(tmp_dir, "word", "settings.xml")
        with open(sx_path, "r", encoding="utf-8") as f:
            sx = f.read()
        sx = re.sub(r'<w:zoom\s+w:val="([^"]+)"/>',
                    r'<w:zoom w:percent="100" w:val="\1"/>', sx)
        with open(sx_path, "w", encoding="utf-8") as f:
            f.write(sx)
        tmp_out = path + ".tmp"
        with zipfile.ZipFile(tmp_out, "w", zipfile.ZIP_DEFLATED) as zout:
            for root, _, files in os.walk(tmp_dir):
                for file in files:
                    fp = os.path.join(root, file)
                    zout.write(fp, os.path.relpath(fp, tmp_dir))
        os.replace(tmp_out, path)
    finally:
        shutil.rmtree(tmp_dir, ignore_errors=True)


# ── main builder ──────────────────────────────────────────────────────────────

def build_docx(data: dict, output_path: str):
    doc = Document()

    # Page: US Letter, 0.75" margins
    sec = doc.sections[0]
    sec.page_width  = Inches(8.5); sec.page_height = Inches(11)
    sec.left_margin = sec.right_margin = sec.top_margin = sec.bottom_margin = Inches(0.75)

    # Default style
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    p   = data.get("personal", {})
    s   = data.get("summary", {}).get("text", "")
    exp = data.get("experience", [])
    edu = data.get("education", [])
    sk  = data.get("skills", [])
    prj = data.get("projects", [])
    cer = data.get("certifications", [])

    # ── NAME ─────────────────────────────────────────────────────────────────
    name_p = _para(doc, space_before=0, space_after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
    nr = name_p.add_run((p.get("name") or "Your Name").upper())
    nr.bold = True; nr.font.name = "Calibri"
    nr.font.size = Pt(20); nr.font.color.rgb = NAVY

    # Job title under name
    if p.get("title"):
        tp = _para(doc, space_before=0, space_after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
        _run(tp, p["title"], bold=True, size=11, color=BLUE)

    # ── CONTACT LINE ─────────────────────────────────────────────────────────
    contacts = [x for x in [
        p.get("email"), p.get("phone"), p.get("location"),
        p.get("linkedin"), p.get("github"), p.get("website")
    ] if x and x.strip()]
    if contacts:
        cp = _para(doc, space_before=0, space_after=5, align=WD_ALIGN_PARAGRAPH.CENTER)
        _run(cp, "  |  ".join(contacts), size=8.5, color=GREY)

    # Gradient-style divider (just a horizontal line via paragraph border)
    div = _para(doc, space_before=0, space_after=0)
    dpPr = div._p.get_or_add_pPr()
    dBdr = OxmlElement("w:pBdr")
    dBot = OxmlElement("w:bottom")
    dBot.set(qn("w:val"), "single"); dBot.set(qn("w:sz"), "12")
    dBot.set(qn("w:space"), "1");    dBot.set(qn("w:color"), "3B82F6")
    dBdr.append(dBot); dpPr.insert(0, dBdr)

    # ── SUMMARY ──────────────────────────────────────────────────────────────
    if s.strip():
        _section_heading(doc, "Professional Summary")
        sp = _para(doc, space_before=2, space_after=2)
        _run(sp, s.strip(), size=9.5, color=GREY)

    # ── EXPERIENCE ───────────────────────────────────────────────────────────
    valid_exp = [e for e in exp if e.get("company") or e.get("role")]
    if valid_exp:
        _section_heading(doc, "Work Experience")
        for e in valid_exp:
            start = _fmt_month(e.get("start",""))
            end   = "Present" if e.get("current") else _fmt_month(e.get("end",""))
            date  = f"{start} – {end}" if start else end
            role_co = f'{e.get("role","")}  —  {e.get("company","")}'
            _two_col(doc, role_co, date, left_bold=True, left_size=10)
            for b in _parse_bullets(e.get("bullets","")):
                _bullet(doc, b)
            _para(doc, space_after=2)

    # ── PROJECTS ─────────────────────────────────────────────────────────────
    valid_prj = [p for p in prj if p.get("name")]
    if valid_prj:
        _section_heading(doc, "Projects")
        for pr in valid_prj:
            title = pr["name"]
            if pr.get("tech"): title += f'  |  {pr["tech"]}'
            tp2 = _para(doc, space_before=3, space_after=1)
            _run(tp2, title, bold=True, size=10)
            if pr.get("url"):
                up = _para(doc, space_before=0, space_after=1)
                _run(up, pr["url"], size=8.5, color=BLUE)
            for b in _parse_bullets(pr.get("bullets","")):
                _bullet(doc, b)
            _para(doc, space_after=2)

    # ── EDUCATION ────────────────────────────────────────────────────────────
    valid_edu = [e for e in edu if e.get("institution") or e.get("degree")]
    if valid_edu:
        _section_heading(doc, "Education")
        for e in valid_edu:
            deg = " in ".join(filter(None, [e.get("degree"), e.get("field")]))
            date = f'{e.get("start","")}{"–" if e.get("start") and e.get("end") else ""}{e.get("end","")}'
            _two_col(doc, deg or "Degree", date, left_bold=True, left_size=10)
            if e.get("institution"):
                ip = _para(doc, space_before=0, space_after=1)
                r = ip.add_run(e["institution"])
                r.italic = True; r.font.name = "Calibri"
                r.font.size = Pt(9.5); r.font.color.rgb = GREY
            if e.get("gpa"):
                gp = _para(doc, space_before=0, space_after=3)
                _run(gp, f'CGPA / Score: {e["gpa"]}', size=9, color=GREY)

    # ── SKILLS ───────────────────────────────────────────────────────────────
    if sk:
        _section_heading(doc, "Skills")
        sp2 = _para(doc, space_before=2, space_after=2)
        _run(sp2, "  •  ".join(sk), size=9.5)

    # ── CERTIFICATIONS ────────────────────────────────────────────────────────
    valid_cer = [c for c in cer if c.get("name")]
    if valid_cer:
        _section_heading(doc, "Certifications")
        for c in valid_cer:
            date = _fmt_month(c.get("date",""))
            _two_col(doc, c["name"], date, left_bold=True, left_size=10)
            if c.get("issuer"):
                ip2 = _para(doc, space_before=0, space_after=3)
                r2 = ip2.add_run(c["issuer"])
                r2.italic = True; r2.font.name = "Calibri"
                r2.font.size = Pt(9); r2.font.color.rgb = GREY

    doc.save(output_path)
    _fix_docx(output_path)
    return output_path
