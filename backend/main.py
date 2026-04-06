"""
ResumeBuilder Backend — NO API KEY NEEDED
Only does: PDF export + DOCX export
Run: uvicorn main:app --reload --port 5000
"""

import io
import os
import uuid
import tempfile
import re
import pypdf
from fastapi import FastAPI, File, UploadFile, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from typing import List

from resume_docx import build_docx
from resume_pdf  import build_pdf

app = FastAPI(title="ResumeBuilder Backend", version="1.0.0")
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ── Models ────────────────────────────────────────────────────────────────────

class Personal(BaseModel):
    name:str=""; title:str=""; email:str=""; phone:str=""
    location:str=""; linkedin:str=""; github:str=""; website:str=""

class Summary(BaseModel):
    text:str=""

class Experience(BaseModel):
    id:str=""; company:str=""; role:str=""
    start:str=""; end:str=""; current:bool=False; bullets: List[str] = []

class Education(BaseModel):
    id:str=""; institution:str=""; degree:str=""
    field:str=""; start:str=""; end:str=""; gpa:str=""

class Project(BaseModel):
    id:str=""; name:str=""; tech:str=""; url:str=""; bullets: List[str] = []

class Certification(BaseModel):
    id:str=""; name:str=""; issuer:str=""; date:str=""

class ResumeData(BaseModel):
    personal:Personal=Personal()
    summary:Summary=Summary()
    experience:List[Experience]=[]
    education:List[Education]=[]
    skills:List[str]=[]
    projects:List[Project]=[]
    certifications:List[Certification]=[]
    templateId: str = "classic" 


# ── Endpoints ─────────────────────────────────────────────────────────────────

@app.get("/")
def root():
    return {"status": "ResumeBuilder backend is running", "version": "1.0.0"}

@app.post("/api/export/pdf")
async def export_pdf(data: ResumeData):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    from reportlab.lib.units import mm

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4,
                            leftMargin=18*mm, rightMargin=18*mm,
                            topMargin=16*mm, bottomMargin=16*mm)

    styles = getSampleStyleSheet()
    story  = []
    tid    = data.templateId

    def fmt_date(s):
        if not s: return ""
        try:
            from datetime import datetime
            return datetime.strptime(s, "%d-%m-%Y").strftime("%d %b %Y")
        except:
            return s

    p = data.personal

    # ── COLOUR PALETTE per template ──
    PALETTES = {
        "classic":   {"accent": colors.HexColor("#0a0a0a"), "sub": colors.HexColor("#555555"), "bullet": "–"},
        "modern":    {"accent": colors.HexColor("#2563eb"), "sub": colors.HexColor("#475569"), "bullet": "▸"},
        "tech":      {"accent": colors.HexColor("#16a34a"), "sub": colors.HexColor("#334155"), "bullet": "→"},
        "executive": {"accent": colors.HexColor("#1e3a5f"), "sub": colors.HexColor("#b8860b"), "bullet": "◆"},
        "etherx":    {"accent": colors.HexColor("#A07830"), "sub": colors.HexColor("#C9A84C"), "bullet": "◆"},
    }
    pal = PALETTES.get(tid, PALETTES["classic"])

    def style(name, **kw):
        return ParagraphStyle(name, parent=styles["Normal"], **kw)

    # Styles
    name_align  = 1 if tid in ("classic",) else 0
    name_s      = style("Name",    fontSize=22, fontName="Helvetica-Bold",
                        alignment=name_align, spaceAfter=2,
                        textColor=colors.HexColor("#0a0a0a"))
    title_s     = style("Title",   fontSize=10, alignment=name_align,
                        spaceAfter=3, textColor=pal["sub"])
    contact_s   = style("Contact", fontSize=8,  alignment=name_align,
                        spaceAfter=10, textColor=colors.HexColor("#555555"))
    sec_s       = style("Sec",     fontSize=10, fontName="Helvetica-Bold",
                        spaceBefore=10, spaceAfter=3, textColor=pal["accent"])
    body_s      = style("Body",    fontSize=9,  spaceAfter=2, leading=13,
                        textColor=colors.HexColor("#1a1a1a"))
    sub_s       = style("Sub",     fontSize=8,  spaceAfter=2,
                        textColor=pal["sub"])
    bullet_s    = style("Bullet",  fontSize=9,  leftIndent=12,
                        spaceAfter=2, leading=13,
                        textColor=colors.HexColor("#1a1a1a"))

    # ── HEADER ──
    story.append(Paragraph(p.name or "Your Name", name_s))
    if p.title:
        story.append(Paragraph(p.title, title_s))
    contacts = "  |  ".join(filter(None, [
        p.email, p.phone, p.location, p.linkedin, p.github, p.website
    ]))
    if contacts:
        story.append(Paragraph(contacts, contact_s))

    # Divider line in accent color
    story.append(HRFlowable(width="100%", thickness=1.5,
                             color=pal["accent"]))
    story.append(Spacer(1, 4))

    def section(title):
        story.append(Spacer(1, 6))
        story.append(Paragraph(title, sec_s))
        story.append(HRFlowable(width="100%", thickness=0.5,
                                 color=pal["sub"]))
        story.append(Spacer(1, 3))

    def bullets(items):
        if not items: return
        # Handle both string (legacy) and list (analysis)
        if isinstance(items, str):
            items = items.split("\n")
        for b in items:
            b = b.strip().lstrip("-•▸◆→*").strip()
            if b:
                story.append(Paragraph(
                    f'<font color="#{pal["accent"].hexval()[2:]}">{pal["bullet"]}</font>  {b}',
                    bullet_s
                ))

    # ── SUMMARY ──
    if data.summary.text:
        section("Professional Summary")
        story.append(Paragraph(data.summary.text, body_s))

    # ── EXPERIENCE ──
    if data.experience:
        section("Work Experience")
        for e in data.experience:
            date_str = f"{fmt_date(e.start)} – {'Present' if e.current else fmt_date(e.end)}"
            story.append(Paragraph(
                f'<b>{e.role}</b>  <font color="#{pal["sub"].hexval()[2:]}">{e.company}</font>'
                f'  <font size="8" color="#888888">({date_str})</font>',
                body_s
            ))
            bullets(e.bullets)
            story.append(Spacer(1, 4))

    # ── PROJECTS ──
    if data.projects:
        section("Projects")
        for pr in data.projects:
            story.append(Paragraph(
                f'<b>{pr.name}</b>  <font color="#{pal["sub"].hexval()[2:]}" size="8">{pr.tech}</font>',
                body_s
            ))
            bullets(pr.bullets)
            story.append(Spacer(1, 4))

    # ── EDUCATION ──
    if data.education:
        section("Education")
        for e in data.education:
            deg = " in ".join(filter(None, [e.degree, e.field]))
            date_str = f"{e.start}{' – ' if e.start and e.end else ''}{e.end}"
            story.append(Paragraph(
                f'<b>{deg}</b>  <font color="#{pal["sub"].hexval()[2:]}">{e.institution}</font>'
                f'  <font size="8" color="#888888">({date_str})</font>',
                body_s
            ))
            if e.gpa:
                story.append(Paragraph(f"CGPA / Score: {e.gpa}", sub_s))
            story.append(Spacer(1, 4))

    # ── SKILLS ──
    if data.skills:
        section("Skills")
        story.append(Paragraph(
            "  •  ".join(data.skills), body_s
        ))

    # ── CERTIFICATIONS ──
    if data.certifications:
        section("Certifications")
        for c in data.certifications:
            story.append(Paragraph(
                f'<b>{c.name}</b>  <font color="#{pal["sub"].hexval()[2:]}">{c.issuer}</font>'
                f'  <font size="8" color="#888888">{fmt_date(c.date)}</font>',
                body_s
            ))
            story.append(Spacer(1, 3))

    doc.build(story)
    buffer.seek(0)
    return StreamingResponse(
        buffer, media_type="application/pdf",
        headers={"Content-Disposition":
                 f'attachment; filename="{p.name or "resume"}_{tid}_resume.pdf"'}
    )

@app.post("/api/export/docx")
async def export_docx(data: ResumeData):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    buffer = io.BytesIO()
    doc    = Document()
    tid    = data.templateId

    # ── Page margins ──
    for section in doc.sections:
        section.top_margin    = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin   = Inches(0.75)
        section.right_margin  = Inches(0.75)

    # ── Colour palettes per template ──
    PALETTES = {
        "classic":   {"accent": (10,10,10),      "sub": (85,85,85),    "bullet": "–",  "center": True},
        "modern":    {"accent": (37,99,235),     "sub": (71,85,105),   "bullet": "▸",  "center": False},
        "tech":      {"accent": (22,163,74),     "sub": (51,65,85),    "bullet": "→",  "center": False},
        "executive": {"accent": (30,58,95),      "sub": (184,134,11),  "bullet": "◆",  "center": False},
        "etherx":    {"accent": (160,120,48),    "sub": (201,168,76),  "bullet": "◆",  "center": False},
    }
    pal = PALETTES.get(tid, PALETTES["classic"])

    def rgb(t): return RGBColor(t[0], t[1], t[2])

    def fmt_date(s):
        if not s: return ""
        try:
            from datetime import datetime
            return datetime.strptime(s, "%d-%m-%Y").strftime("%d %b %Y")
        except:
            return s

    def add_hr(color_tuple=(180,180,180)):
        """Add a horizontal rule paragraph"""
        p   = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr= OxmlElement("w:pBdr")
        bot = OxmlElement("w:bottom")
        bot.set(qn("w:val"),   "single")
        bot.set(qn("w:sz"),    "6")
        bot.set(qn("w:space"), "1")
        bot.set(qn("w:color"),
                "{:02X}{:02X}{:02X}".format(*color_tuple))
        pBdr.append(bot)
        pPr.append(pBdr)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(2)
        return p

    def add_section_heading(title):
        p   = doc.add_paragraph()
        run = p.add_run(title.upper())
        run.bold           = True
        run.font.size      = Pt(10)
        run.font.color.rgb = rgb(pal["accent"])
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(2)
        add_hr(pal["sub"])

    def add_bullet(text):
        text = text.strip().lstrip("-•▸◆→*").strip()
        if not text: return
        p   = doc.add_paragraph()
        bul = p.add_run(f'{pal["bullet"]}  ')
        bul.font.color.rgb = rgb(pal["accent"])
        bul.font.size      = Pt(9)
        run = p.add_run(text)
        run.font.size      = Pt(9)
        p.paragraph_format.left_indent  = Inches(0.15)
        p.paragraph_format.space_after  = Pt(1)

    def add_bullets(items):
        if not items: return
        # Handle both string (legacy) and list (analysis)
        if isinstance(items, str):
            items = items.split("\n")
        for b in items:
            b = b.strip()
            if b: add_bullet(b)

    p = data.personal

    # ── NAME ──
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER if pal["center"] else WD_ALIGN_PARAGRAPH.LEFT
    name_r = name_p.add_run(p.name or "Your Name")
    name_r.bold            = True
    name_r.font.size       = Pt(22)
    name_r.font.color.rgb  = RGBColor(10,10,10)
    name_p.paragraph_format.space_after = Pt(2)

    # ── TITLE ──
    if p.title:
        tp = doc.add_paragraph()
        tp.alignment = WD_ALIGN_PARAGRAPH.CENTER if pal["center"] else WD_ALIGN_PARAGRAPH.LEFT
        tr = tp.add_run(p.title)
        tr.font.size      = Pt(10)
        tr.font.color.rgb = rgb(pal["sub"])
        tp.paragraph_format.space_after = Pt(2)

    # ── CONTACTS ──
    contacts = "  |  ".join(filter(None, [
        p.email, p.phone, p.location, p.linkedin, p.github, p.website
    ]))
    if contacts:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER if pal["center"] else WD_ALIGN_PARAGRAPH.LEFT
        cr = cp.add_run(contacts)
        cr.font.size      = Pt(8)
        cr.font.color.rgb = RGBColor(85,85,85)
        cp.paragraph_format.space_after = Pt(4)

    # Accent HR under header
    add_hr(pal["accent"])

    # ── SUMMARY ──
    if data.summary.text:
        add_section_heading("Professional Summary")
        sp = doc.add_paragraph()
        sr = sp.add_run(data.summary.text)
        sr.font.size = Pt(9)
        sp.paragraph_format.space_after = Pt(4)

    # ── EXPERIENCE ──
    if data.experience:
        add_section_heading("Work Experience")
        for e in data.experience:
            ep  = doc.add_paragraph()
            r1  = ep.add_run(e.role or "Role")
            r1.bold = True; r1.font.size = Pt(10)
            r2  = ep.add_run(f"  —  {e.company}")
            r2.font.size      = Pt(10)
            r2.font.color.rgb = rgb(pal["sub"])
            date_str = f"{fmt_date(e.start)} – {'Present' if e.current else fmt_date(e.end)}"
            r3  = ep.add_run(f"  ({date_str})")
            r3.font.size      = Pt(8)
            r3.font.color.rgb = RGBColor(136,136,136)
            ep.paragraph_format.space_before = Pt(6)
            ep.paragraph_format.space_after  = Pt(2)
            add_bullets(e.bullets)

    # ── PROJECTS ──
    if data.projects:
        add_section_heading("Projects")
        for pr in data.projects:
            pp2 = doc.add_paragraph()
            r1  = pp2.add_run(pr.name or "Project")
            r1.bold = True; r1.font.size = Pt(10)
            if pr.tech:
                r2 = pp2.add_run(f"  |  {pr.tech}")
                r2.font.size      = Pt(9)
                r2.font.color.rgb = rgb(pal["sub"])
            pp2.paragraph_format.space_before = Pt(6)
            pp2.paragraph_format.space_after  = Pt(2)
            add_bullets(pr.bullets)

    # ── EDUCATION ──
    if data.education:
        add_section_heading("Education")
        for e in data.education:
            deg  = " in ".join(filter(None, [e.degree, e.field]))
            date_str = f"{e.start}{' – ' if e.start and e.end else ''}{e.end}"
            ep   = doc.add_paragraph()
            r1   = ep.add_run(deg or "Degree")
            r1.bold = True; r1.font.size = Pt(10)
            r2   = ep.add_run(f"  —  {e.institution}")
            r2.font.size      = Pt(10)
            r2.font.color.rgb = rgb(pal["sub"])
            r3   = ep.add_run(f"  ({date_str})")
            r3.font.size      = Pt(8)
            r3.font.color.rgb = RGBColor(136,136,136)
            ep.paragraph_format.space_before = Pt(6)
            ep.paragraph_format.space_after  = Pt(2)
            if e.gpa:
                gp  = doc.add_paragraph()
                gr  = gp.add_run(f"CGPA / Score: {e.gpa}")
                gr.font.size      = Pt(9)
                gr.font.color.rgb = rgb(pal["sub"])
                gp.paragraph_format.space_after = Pt(2)

    # ── SKILLS ──
    if data.skills:
        add_section_heading("Skills")
        skp = doc.add_paragraph()
        skr = skp.add_run("  •  ".join(data.skills))
        skr.font.size = Pt(9)
        skp.paragraph_format.space_after = Pt(4)

    # ── CERTIFICATIONS ──
    if data.certifications:
        add_section_heading("Certifications")
        for c in data.certifications:
            cp2 = doc.add_paragraph()
            r1  = cp2.add_run(c.name)
            r1.bold = True; r1.font.size = Pt(10)
            if c.issuer:
                r2 = cp2.add_run(f"  —  {c.issuer}")
                r2.font.size      = Pt(9)
                r2.font.color.rgb = rgb(pal["sub"])
            if c.date:
                r3 = cp2.add_run(f"  {fmt_date(c.date)}")
                r3.font.size      = Pt(8)
                r3.font.color.rgb = RGBColor(136,136,136)
            cp2.paragraph_format.space_before = Pt(4)
            cp2.paragraph_format.space_after  = Pt(2)

    doc.save(buffer)
    buffer.seek(0)
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition":
                 f'attachment; filename="{p.name or "resume"}_{tid}_resume.docx"'}
    )

@app.post("/api/analyze")
async def analyze_resume(
    role: str = Form(""),
    job_description: str = Form(""),
    file: UploadFile = File(...)
):
    try:
        text = ""
        if file.filename.lower().endswith(".pdf"):
            reader = pypdf.PdfReader(file.file)
            for page in reader.pages:
                text += page.extract_text() + "\n"
        elif file.filename.lower().endswith(".docx"):
            from docx import Document
            doc = Document(file.file)
            text = "\n".join([para.text for para in doc.paragraphs])
        else:
            raise HTTPException(status_code=400, detail="Invalid file format")
        text = text.lower()
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Could not parse file: {str(e)}")

    STOPWORDS = {"a", "about", "above", "after", "again", "against", "all", "am", "an", "and", "any", "are", "aren't", "as", "at", "be", "because", "been", "before", "being", "below", "between", "both", "but", "by", "can't", "cannot", "could", "couldn't", "did", "didn't", "do", "does", "doesn't", "doing", "don't", "down", "during", "each", "few", "for", "from", "further", "had", "hadn't", "has", "hasn't", "have", "haven't", "having", "he", "he'd", "he'll", "he's", "her", "here", "here's", "hers", "herself", "him", "himself", "his", "how", "how's", "i", "i'd", "i'll", "i'm", "i've", "if", "in", "into", "is", "isn't", "it", "it's", "its", "itself", "let's", "me", "more", "most", "mustn't", "my", "myself", "no", "nor", "not", "of", "off", "on", "once", "only", "or", "other", "ought", "our", "ours", "ourselves", "out", "over", "own", "same", "shan't", "she", "she'd", "she'll", "she's", "should", "shouldn't", "so", "some", "such", "than", "that", "that's", "the", "their", "theirs", "them", "themselves", "then", "there", "there's", "these", "they", "they'd", "they'll", "they're", "they've", "this", "those", "through", "to", "too", "under", "until", "up", "very", "was", "wasn't", "we", "we'd", "we'll", "we're", "we've", "were", "weren't", "what", "what's", "when", "when's", "where", "where's", "which", "while", "who", "who's", "whom", "why", "why's", "with", "won't", "would", "wouldn't", "you", "you'd", "you'll", "you're", "you've", "your", "yours", "yourself", "yourselves", "experience", "role", "team", "work", "skills", "knowledge", "years", "working"}

    # Extract keywords from JD
    jd_words = re.findall(r'[a-zA-Z0-9+#]+', job_description.lower())
    potential_keywords = set()
    for w in jd_words:
        if len(w) > 2 and w not in STOPWORDS and not w.isnumeric():
            potential_keywords.add(w)

    matched_keywords = []
    missing_keywords = []
    
    for kw in potential_keywords:
        if re.search(rf'\b{re.escape(kw)}\b', text):
            matched_keywords.append(kw)
        elif kw in text:
            matched_keywords.append(kw)
        else:
            missing_keywords.append(kw)

    total_kw = len(potential_keywords)
    matched_count = len(matched_keywords)
    
    keyword_match_score = (matched_count / total_kw * 100) if total_kw > 0 else 0
    skill_alignment_score = min(keyword_match_score * 1.1, 100)
    
    action_verbs = ["developed", "managed", "led", "created", "built", "designed", "improved", "increased", "reduced", "delivered", "implemented", "achieved"]
    verbs_found = sum(1 for v in action_verbs if re.search(rf'\b{v}\b', text))
    experience_relevance_score = min(50 + (verbs_found * 10), 100)
    
    headers = ["experience", "education", "skills", "projects", "summary"]
    headers_found = sum(1 for h in headers if re.search(rf'\b{h}\b', text))
    ats_compliance_score = (headers_found / len(headers)) * 100
    
    if total_kw == 0:
        overall_score = (experience_relevance_score * 0.5) + (ats_compliance_score * 0.5)
    else:
        overall_score = (keyword_match_score * 0.4) + (skill_alignment_score * 0.3) + (experience_relevance_score * 0.2) + (ats_compliance_score * 0.1)
        
    improvement_suggestions = []
    if missing_keywords:
        improvement_suggestions.append(f"Consider adding missing keywords like: {', '.join(list(missing_keywords)[:5])}")
    if verbs_found < 3:
        improvement_suggestions.append("Use more action verbs (e.g. Led, Developed, Improved) to describe experiences.")
    if headers_found < 3:
        improvement_suggestions.append("Ensure your resume has standard sections like Experience, Education, and Skills.")
    if not re.search(r'\b\d+%\b|\b\$\d+\b', text):
        improvement_suggestions.append("Add quantifiable achievements (e.g. increased sales by 20%, saved $10k).")
        
    return {
        "overall_score": overall_score,
        "keyword_match_score": keyword_match_score,
        "matched_keywords": matched_keywords,
        "missing_keywords": missing_keywords,
        "skill_alignment_score": skill_alignment_score,
        "experience_relevance_score": experience_relevance_score,
        "ats_compliance_score": ats_compliance_score,
        "improvement_suggestions": improvement_suggestions
    }

# uvicorn main:app --reload --port 5000
